from docx import Document

path = r"e:\Web\WLP website\مادران و دختران.docx"
doc = Document(path)

out = open(r"C:\Users\DELL\AppData\Local\Temp\claude\e--Web-WLP-website\1cd0673d-e43b-4e73-b486-1cd33e221790\scratchpad\docx_compare.txt", "w", encoding="utf-8")
def p(*a): print(*a, file=out)

# find paragraphs mentioning "اواخر دهه" or "درآمد" near the شاه‌کار/تابوت‌های رویین section
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "اواخر" in text and "۱۹۸۰" in text:
        p(f"[{i}] MATCH (aвакhar+1980): {text}")
    if text == "درآمد":
        p(f"[{i}] HEADING درآمد, style={para.style.name}")

p("\n=== paragraphs 394-420 (شاه‌کار section start) for context ===")
for i in range(394, 420):
    t = doc.paragraphs[i].text.strip()
    if t:
        p(f"[{i}] {t}")
